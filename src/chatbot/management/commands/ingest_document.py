from pathlib import Path

from django.core.management.base import BaseCommand, CommandError
from django.db import transaction

from chatbot.infrastructure.embeddings.openai_embedder import OpenAIEmbedder
from chatbot.models import Chunk


def chunk_text(text: str, max_chars: int = 1000) -> list[str]:
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for paragraph in paragraphs:
        if current_len + len(paragraph) + 1 > max_chars and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(paragraph)
        current_len += len(paragraph) + 1

    if current:
        chunks.append("\n".join(current))

    return chunks


def load_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise CommandError("pypdf is required for PDF ingestion. Run: pip install pypdf") from exc

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise CommandError("python-docx is required for DOCX ingestion. Run: pip install python-docx") from exc

        doc = Document(str(path))
        parts = []
        parts.extend([p.text for p in doc.paragraphs if p.text.strip()])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend([p.text for p in cell.paragraphs if p.text.strip()])

        for section in doc.sections:
            header = section.header
            if header:
                parts.extend([p.text for p in header.paragraphs if p.text.strip()])
            footer = section.footer
            if footer:
                parts.extend([p.text for p in footer.paragraphs if p.text.strip()])

        text = "\n".join(parts).strip()
        if text:
            return text

        try:
            import docx2txt
        except ImportError as exc:
            raise CommandError(
                "DOCX appears to use text boxes. Install docx2txt: pip install docx2txt"
            ) from exc

        return docx2txt.process(str(path)).strip()

    return path.read_text(encoding="utf-8")


class Command(BaseCommand):
    help = "Ingest a text document into the vector store."

    def add_arguments(self, parser):
        parser.add_argument("path", type=str, help="Path to a text file.")
        parser.add_argument("--source", type=str, default=None, help="Source label.")
        parser.add_argument(
            "--clear",
            action="store_true",
            help="Delete existing chunks before ingesting.",
        )

    def handle(self, *args, **options):
        path = Path(options["path"])
        if not path.exists():
            raise CommandError(f"File not found: {path}")

        source = options["source"] or path.name
        text = load_text(path)
        chunks = chunk_text(text)
        if not chunks:
            raise CommandError("No content to ingest.")

        embedder = OpenAIEmbedder()

        with transaction.atomic():
            if options["clear"]:
                Chunk.objects.all().delete()

            objects = []
            for chunk in chunks:
                embedding = embedder.embed(chunk)
                objects.append(
                    Chunk(
                        content=chunk,
                        source=source,
                        embedding=embedding,
                    )
                )

            Chunk.objects.bulk_create(objects)

        self.stdout.write(self.style.SUCCESS(f"Ingested {len(objects)} chunks."))
